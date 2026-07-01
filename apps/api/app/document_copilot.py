from __future__ import annotations

import json
import uuid
from io import BytesIO
from typing import Any

from openai import OpenAI
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, FieldCondition, Filter, MatchValue, PointStruct, VectorParams

from app.config import get_settings
from app.memory import embed_text
from app.models import Document

DOCUMENT_EMBEDDING_SIZE = 1536
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def _openai_client() -> OpenAI | None:
    settings = get_settings()
    if not settings.openai_api_key or settings.openai_api_key in {"replace-me", "sk-xxxxxxxx"}:
        return None
    return OpenAI(api_key=settings.openai_api_key)


def _qdrant_client() -> QdrantClient | None:
    settings = get_settings()
    if not settings.qdrant_url:
        return None
    return QdrantClient(url=settings.qdrant_url, api_key=settings.qdrant_api_key or None, timeout=5)


def _ensure_document_collection(client: QdrantClient) -> None:
    settings = get_settings()
    collections = client.get_collections().collections
    if any(collection.name == settings.qdrant_document_collection for collection in collections):
        return
    client.create_collection(
        collection_name=settings.qdrant_document_collection,
        vectors_config=VectorParams(size=DOCUMENT_EMBEDDING_SIZE, distance=Distance.COSINE),
    )


def extract_text(file_name: str, mime_type: str, data: bytes) -> tuple[str, dict[str, Any]]:
    suffix = "." + file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError("Only PDF, DOCX, and TXT files are supported")

    if suffix == ".pdf" or mime_type == "application/pdf":
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(data))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
        text = "\n\n".join(page for page in pages if page)
        return normalize_text(text), {"page_count": len(reader.pages), "file_type": "pdf"}

    if suffix == ".docx" or mime_type.endswith("wordprocessingml.document"):
        from docx import Document as DocxDocument

        doc = DocxDocument(BytesIO(data))
        paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        table_lines: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_lines.append(" | ".join(cells))
        text = "\n".join([*paragraphs, *table_lines])
        return normalize_text(text), {"paragraph_count": len(paragraphs), "file_type": "docx"}

    text = data.decode("utf-8", errors="ignore")
    return normalize_text(text), {"file_type": "txt"}


def normalize_text(value: str) -> str:
    lines = [line.strip() for line in value.replace("\r", "\n").split("\n")]
    compact: list[str] = []
    previous_blank = False
    for line in lines:
        blank = not line
        if blank and previous_blank:
            continue
        compact.append(line)
        previous_blank = blank
    return "\n".join(compact).strip()


def chunk_text(text: str, chunk_size: int = 1200, overlap: int = 160) -> list[str]:
    text = normalize_text(text)
    if not text:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + chunk_size)
        chunks.append(text[start:end].strip())
        if end == len(text):
            break
        start = max(0, end - overlap)
    return [chunk for chunk in chunks if chunk]


def upsert_document_embeddings(document: Document) -> list[str]:
    chunks = chunk_text(document.extracted_text)
    point_ids: list[str] = []
    if not chunks:
        return point_ids
    client = _qdrant_client()
    if client is None:
        return point_ids
    try:
        _ensure_document_collection(client)
        points: list[PointStruct] = []
        for index, chunk in enumerate(chunks[:200]):
            vector = embed_text(chunk)
            if vector is None:
                continue
            point_id = str(uuid.uuid4())
            point_ids.append(point_id)
            points.append(
                PointStruct(
                    id=point_id,
                    vector=vector,
                    payload={
                        "document_id": str(document.id),
                        "user_id": str(document.user_id),
                        "chunk_index": index,
                        "title": document.title,
                        "file_name": document.file_name,
                        "text": chunk,
                    },
                )
            )
        if points:
            client.upsert(collection_name=get_settings().qdrant_document_collection, points=points)
    except Exception:
        return point_ids
    return point_ids


def retrieve_document_context(document: Document, question: str, limit: int = 6) -> list[str]:
    vector = embed_text(question)
    if vector is not None:
        client = _qdrant_client()
        if client is not None:
            try:
                _ensure_document_collection(client)
                hits = client.search(
                    collection_name=get_settings().qdrant_document_collection,
                    query_vector=vector,
                    query_filter=Filter(
                        must=[
                            FieldCondition(key="user_id", match=MatchValue(value=str(document.user_id))),
                            FieldCondition(key="document_id", match=MatchValue(value=str(document.id))),
                        ]
                    ),
                    limit=limit,
                )
                contexts = [str(hit.payload.get("text", "")) for hit in hits if hit.payload and hit.payload.get("text")]
                if contexts:
                    return contexts
            except Exception:
                pass
    return chunk_text(document.extracted_text)[:limit]


def generate_document_copilot(document: Document, action: str, question: str = "") -> dict[str, Any]:
    contexts = retrieve_document_context(document, question or action) if action == "question" else chunk_text(document.extracted_text)[:8]
    content = "\n\n".join(contexts)[:14000]
    client = _openai_client()
    if client is None:
        return _fallback_document_response(document, action, question, content)

    instruction = {
        "summary": "Summarize the document with a concise executive summary and important points.",
        "question": "Answer the user's question using only the document context. Say when the answer is not present.",
        "notes": "Generate clean study/work notes from the document with headings and bullets.",
        "action_items": "Extract concrete action items, owners if mentioned, and suggested next steps.",
    }.get(action, "Analyze the document.")
    prompt = {
        "document_title": document.title,
        "file_name": document.file_name,
        "action": action,
        "question": question,
        "document_context": content,
        "instructions": instruction + " Return JSON with title, answer, and items array. Items must have title and description.",
    }
    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": "You are HumanOS Document Copilot. Be precise, source-grounded, and practical."},
                {"role": "user", "content": json.dumps(prompt)},
            ],
            temperature=0.25,
            response_format={"type": "json_object"},
        )
        payload = json.loads(response.choices[0].message.content or "{}")
        return {
            "action": action,
            "title": str(payload.get("title") or _title_for(action)),
            "answer": str(payload.get("answer") or ""),
            "items": _normalize_items(payload.get("items")),
        }
    except Exception:
        return _fallback_document_response(document, action, question, content)


def _fallback_document_response(document: Document, action: str, question: str, content: str) -> dict[str, Any]:
    sentences = [part.strip() for part in content.replace("\n", " ").split(".") if len(part.strip()) > 24]
    preview = ". ".join(sentences[:4])[:1200]
    if action == "question":
        keywords = [word.lower() for word in question.split() if len(word) > 3]
        matches = [sentence for sentence in sentences if any(keyword in sentence.lower() for keyword in keywords)]
        answer = ". ".join(matches[:4]) or "I could not find a direct answer in the extracted document text."
        return {"action": action, "title": "Document answer", "answer": answer, "items": []}
    if action == "notes":
        return {
            "action": action,
            "title": "Generated notes",
            "answer": preview or "No readable document text was extracted.",
            "items": [{"title": f"Note {index + 1}", "description": sentence[:240]} for index, sentence in enumerate(sentences[:6])],
        }
    if action == "action_items":
        markers = ["must", "should", "need", "action", "next", "follow", "deadline", "due"]
        action_sentences = [sentence for sentence in sentences if any(marker in sentence.lower() for marker in markers)] or sentences[:4]
        return {
            "action": action,
            "title": "Action items",
            "answer": f"Extracted action candidates from {document.title}.",
            "items": [{"title": sentence[:80], "description": sentence[:260]} for sentence in action_sentences[:8]],
        }
    return {
        "action": action,
        "title": "Document summary",
        "answer": preview or "No readable document text was extracted.",
        "items": [{"title": "Key point", "description": sentence[:260]} for sentence in sentences[:5]],
    }


def _title_for(action: str) -> str:
    return {"summary": "Document summary", "question": "Document answer", "notes": "Generated notes", "action_items": "Action items"}.get(action, "Document insight")


def _normalize_items(value: Any) -> list[dict[str, str]]:
    if not isinstance(value, list):
        return []
    items: list[dict[str, str]] = []
    for item in value[:12]:
        if isinstance(item, dict):
            title = str(item.get("title") or "Insight").strip()
            description = str(item.get("description") or item.get("body") or "").strip()
        else:
            title = "Insight"
            description = str(item).strip()
        if title or description:
            items.append({"title": title[:140], "description": description[:500]})
    return items
